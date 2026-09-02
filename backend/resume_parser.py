"""
resume_parser.py
Extracts raw text from an uploaded resume file (PDF, DOCX, or TXT).
"""

import io
import pdfplumber
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of tables, since some resumes use table layouts
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)

    return "\n".join(paragraphs)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Dispatches to the right extractor based on file extension.
    Raises ValueError for unsupported file types.
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif lower_name.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ValueError(
            f"Unsupported file type: {filename}. Please upload a .pdf, .docx, or .txt file."
        )
